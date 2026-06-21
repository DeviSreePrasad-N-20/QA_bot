"""
ingest.py
---------
Ingestion pipeline: scans data/, extracts text from PDFs and DOCX files,
splits the text into overlapping chunks, embeds them with Gemini's
text-embedding-004, and persists everything to a local ChromaDB store.

Run once whenever your source documents change:

    python -m src.ingest
    python -m src.ingest --rebuild   # wipe and re-index from scratch
"""

import argparse
import os
import sys
from pathlib import Path

# Allow this file to be run directly (`python src/ingest.py`) as well as as
# a package module (`python -m src.ingest`) without import errors.
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from pypdf import PdfReader
from docx import Document
from docx.oxml.ns import qn
from tqdm import tqdm

from src import config
from src.embeddings import embed_documents


# ---------------------------------------------------------------------------
# Step 2: Document extraction
# ---------------------------------------------------------------------------

def extract_pdf_pages(file_path: str) -> list[dict]:
    """
    Extracts text page-by-page from a PDF, tracking page numbers and file
    source so every chunk can later be traced back to exactly where it came
    from.
    """
    extracted_data = []
    file_name = os.path.basename(file_path)

    try:
        reader = PdfReader(file_path)
        for index, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                # Collapse repeated whitespace/newlines introduced by PDF
                # layout streams into clean, single-spaced text.
                clean_text = " ".join(text.split())
                extracted_data.append({
                    "text": clean_text,
                    "metadata": {
                        "source": file_name,
                        "page": index + 1,  # 1-indexed for human readability
                    },
                })
    except Exception as e:
        print(f"  [error] Could not read PDF {file_name}: {e}")

    return extracted_data


def extract_docx_pages(file_path: str) -> list[dict]:
    """
    Extracts text from a DOCX file.

    Word documents do not store fixed page numbers the way PDFs do (true
    pagination depends on rendering, fonts, and screen size). We approximate
    "pages" by splitting at explicit page-break runs (Insert > Page Break in
    Word). A document with no explicit page breaks is returned as a single
    logical page. Table content is appended to the page it appears in.
    """
    file_name = os.path.basename(file_path)
    extracted_data = []

    try:
        doc = Document(file_path)
        current_page_parts: list[str] = []
        page_num = 1

        def flush():
            nonlocal page_num
            text = " ".join(" ".join(current_page_parts).split())
            if text:
                extracted_data.append({
                    "text": text,
                    "metadata": {"source": file_name, "page": page_num},
                })
                page_num += 1
            current_page_parts.clear()

        for para in doc.paragraphs:
            has_page_break = any(
                br.get(qn("w:type")) == "page"
                for run in para.runs
                for br in run._element.findall(qn("w:br"))
            )
            if para.text.strip():
                current_page_parts.append(para.text.strip())
            if has_page_break:
                flush()

        # Tables aren't anchored to a specific paragraph position via
        # python-docx, so we append them at the end as supplementary content.
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    current_page_parts.append(row_text)

        flush()  # flush whatever remains (the last/only page)
    except Exception as e:
        print(f"  [error] Could not read DOCX {file_name}: {e}")

    return extracted_data


EXTRACTORS = {
    ".pdf": extract_pdf_pages,
    ".docx": extract_docx_pages,
}


def discover_and_extract(data_dir: Path = config.DATA_DIR) -> list[dict]:
    """Scans data_dir for supported files and extracts page-level text from each."""
    files = sorted(
        f for f in Path(data_dir).iterdir()
        if f.suffix.lower() in EXTRACTORS and f.is_file()
    )

    if not files:
        print(f"No supported documents (.pdf, .docx) found in {data_dir}")
        return []

    all_pages: list[dict] = []
    for file_path in tqdm(files, desc="Extracting documents"):
        extractor = EXTRACTORS[file_path.suffix.lower()]
        pages = extractor(str(file_path))
        all_pages.extend(pages)

    return all_pages


# ---------------------------------------------------------------------------
# Step 3: Recursive text chunking
# ---------------------------------------------------------------------------

# Ordered from "most semantically meaningful" to "last resort": we always try
# to cut on paragraph breaks first, then line breaks, then spaces, and only
# fall back to a hard character cut if a single word is itself longer than
# the chunk size (e.g. a long URL or code token).
_SEPARATORS = ["\n\n", "\n", " ", ""]


def _split_on_separator(text: str, separators: list[str], chunk_size: int) -> list[str]:
    """Recursively splits text, trying coarser separators before finer ones."""
    if not text:
        return []

    sep = separators[0]
    remaining_separators = separators[1:]

    pieces = text.split(sep) if sep else list(text)

    results: list[str] = []
    for piece in pieces:
        if not piece:
            continue
        if len(piece) <= chunk_size or not remaining_separators:
            results.append(piece)
        else:
            # This piece is still too big even after splitting on `sep`;
            # recurse with the next, finer-grained separator.
            results.extend(_split_on_separator(piece, remaining_separators, chunk_size))
    return results


def _merge_with_overlap(pieces: list[str], chunk_size: int, chunk_overlap: int) -> list[str]:
    """
    Greedily packs small pieces back together (joined by a single space)
    into chunks as close to chunk_size as possible, carrying a tail of
    `chunk_overlap` characters from one chunk into the start of the next so
    information sitting on a cut boundary isn't lost.
    """
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0

    for piece in pieces:
        piece_len = len(piece)
        added_len = piece_len + (1 if current else 0)  # +1 for the joining space

        if current and current_len + added_len > chunk_size:
            chunks.append(" ".join(current))

            # Build the overlap tail from the end of the current chunk.
            overlap: list[str] = []
            overlap_len = 0
            for p in reversed(current):
                extra = len(p) + (1 if overlap else 0)
                if overlap_len + extra > chunk_overlap:
                    break
                overlap.insert(0, p)
                overlap_len += extra

            current = overlap
            current_len = overlap_len

        current.append(piece)
        current_len += piece_len + (1 if len(current) > 1 else 0)

    if current:
        chunks.append(" ".join(current))

    return chunks


def recursive_character_split(
    text: str,
    chunk_size: int = config.CHUNK_SIZE,
    chunk_overlap: int = config.CHUNK_OVERLAP,
) -> list[str]:
    """Splits a single block of text into overlapping, semantically-aware chunks."""
    pieces = _split_on_separator(text, _SEPARATORS, chunk_size)
    return _merge_with_overlap(pieces, chunk_size, chunk_overlap)


def chunk_extracted_pages(
    pages: list[dict],
    chunk_size: int = config.CHUNK_SIZE,
    chunk_overlap: int = config.CHUNK_OVERLAP,
) -> list[dict]:
    """
    Splits page-level documents into smaller, overlapping chunks, carrying
    the original source/page metadata over to every individual chunk.
    """
    chunks = []
    for page in pages:
        text_chunks = recursive_character_split(page["text"], chunk_size, chunk_overlap)
        for i, chunk_text in enumerate(text_chunks):
            chunks.append({
                "text": chunk_text,
                "metadata": {**page["metadata"], "chunk_index": i},
            })
    return chunks


# ---------------------------------------------------------------------------
# Step 4: Embedding + persistence
# ---------------------------------------------------------------------------

def _get_collection():
    """Creates (or opens) the persistent ChromaDB collection used for storage."""
    import chromadb

    config.require_api_key()
    config.DB_DIR.mkdir(parents=True, exist_ok=True)

    client = chromadb.PersistentClient(path=str(config.DB_DIR))
    collection = client.get_or_create_collection(
        name=config.COLLECTION_NAME,
        metadata={"hnsw:space": "cosine"},
    )
    return client, collection


def save_to_vector_db(chunks: list[dict], rebuild: bool = False) -> int:
    """
    Embeds text chunks (via Gemini) and saves them into the persistent,
    disk-backed ChromaDB collection. Uploads happen in batches so a single
    large document library doesn't hit per-request payload limits.
    """
    if not chunks:
        print("No chunks to index.")
        return 0

    client, collection = _get_collection()

    if rebuild:
        client.delete_collection(config.COLLECTION_NAME)
        _, collection = _get_collection()

    batch_size = config.EMBED_BATCH_SIZE
    total_indexed = 0

    for start in tqdm(range(0, len(chunks), batch_size), desc="Embedding & indexing"):
        batch = chunks[start:start + batch_size]
        ids = [f"chunk_{start + i}" for i in range(len(batch))]
        documents = [c["text"] for c in batch]
        metadatas = [c["metadata"] for c in batch]

        embeddings = embed_documents(documents)

        collection.add(
            ids=ids,
            documents=documents,
            metadatas=metadatas,
            embeddings=embeddings,
        )
        total_indexed += len(batch)

    return total_indexed


def ensure_vector_db(chunk_size: int = config.CHUNK_SIZE, chunk_overlap: int = config.CHUNK_OVERLAP) -> int:
    """Build the vector database on demand when it doesn't exist yet."""
    if config.DB_DIR.exists() and any(config.DB_DIR.iterdir()):
        return 0

    pages = discover_and_extract()
    if not pages:
        raise RuntimeError(
            f"No supported documents (.pdf, .docx) found in {config.DATA_DIR}"
        )

    chunks = chunk_extracted_pages(pages, chunk_size, chunk_overlap)
    return save_to_vector_db(chunks, rebuild=False)


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Ingest documents into the RAG vector store.")
    parser.add_argument(
        "--rebuild", action="store_true",
        help="Delete the existing collection and re-index from scratch.",
    )
    parser.add_argument(
        "--chunk-size", type=int, default=config.CHUNK_SIZE,
        help=f"Characters per chunk (default: {config.CHUNK_SIZE}).",
    )
    parser.add_argument(
        "--chunk-overlap", type=int, default=config.CHUNK_OVERLAP,
        help=f"Character overlap between chunks (default: {config.CHUNK_OVERLAP}).",
    )
    args = parser.parse_args()

    print(f"Scanning '{config.DATA_DIR}' for documents...")
    pages = discover_and_extract()
    if not pages:
        sys.exit(1)
    print(f"Extracted {len(pages)} page(s) from the source documents.")

    chunks = chunk_extracted_pages(pages, args.chunk_size, args.chunk_overlap)
    print(f"Split into {len(chunks)} chunk(s) "
          f"(chunk_size={args.chunk_size}, overlap={args.chunk_overlap}).")

    try:
        count = save_to_vector_db(chunks, rebuild=args.rebuild)
    except EnvironmentError as e:
        print(f"\n[error] {e}")
        sys.exit(1)

    print(f"\nSuccessfully indexed {count} chunks in '{config.DB_DIR}'.")
    print("You can now run: python -m src.main   (or)   streamlit run src/main.py")


if __name__ == "__main__":
    main()
